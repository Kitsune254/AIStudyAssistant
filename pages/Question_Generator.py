import os
import json
import streamlit as st
import fitz  # PyMuPDF
from typing import List, Dict, Any
from dataclasses import dataclass
from dotenv import load_dotenv

# Optional: save results to docx
from docx import Document

load_dotenv()  # Load .env file
gemini_api_key = os.getenv('GEMINI_API_KEY')

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except Exception:
    GEMINI_AVAILABLE = False

def gemini_configure():
    gemini_api_key = os.getenv('GEMINI_API_KEY')
    if not gemini_api_key:
        raise RuntimeError("Please set the GEMINI_API_KEY environment variable.")
    if not GEMINI_AVAILABLE:
        raise RuntimeError("google.generativeai package not installed. Install it or adapt this function.")
    genai.configure(api_key=gemini_api_key)

def gemini_generate_text(prompt: str, model: str = "text-bison-001", max_output_tokens: int = 1024) -> str:
    """
    Calls Gemini (Generative Language) and returns the text content.
    Replace/modify this if you use a different client or REST.
    """
    gemini_configure()
    # The library may provide 'generate_text' or 'client.generate' methods depending on versions.
    # We'll try a common interface; adapt if your library differs.
    resp = genai.generate_text(model=model, prompt=prompt, max_output_tokens=max_output_tokens)
    # The response object shape can differ; convert to string gracefully
    if hasattr(resp, "text"):
        return resp.text
    # fallback
    return str(resp)

# ---------- Helpers ----------
def extract_text_from_pdf(uploaded_file) -> str:
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    full_text = []
    for page in doc:
        text = page.get_text()
        if text:
            full_text.append(text)
    return "\n\n".join(full_text)

@dataclass
class Question:
    id: int
    qtype: str  # 'mcq' or 'open'
    question: str
    options: List[str]  # only for mcq
    correct_answer: str = None  # grader will use this; don't show to user
    explanation: str = None

def ask_gemini_for_questions(pdf_text: str, num_questions: int = 6) -> List[Question]:
    """
    Prompt Gemini to generate questions in strict JSON:
    [
      {"id":1, "type":"mcq", "question":"...", "options":["A","B","C","D"], "answer":"B", "explanation":"..."},
      ...
    ]
    """
    prompt = f"""
You are an assistant that reads a document and generates educational questions.

Instructions:
- Read the provided document text delimited by <<DOC>> and <<ENDDOC>>.
- Produce exactly {num_questions} questions covering the document's main ideas and details.
- For each question produce either:
  - an MCQ with 4 options (labelled or unlabelled) and the correct option (single letter or full option text), OR
  - an open question where the user must type an answer.
- Output MUST be strict JSON (no extra commentary) with an array of objects:
[
  {{
    "id": 1,
    "type": "mcq" or "open",
    "question": "question text",
    "options": ["opt1","opt2","opt3","opt4"],   // only for mcq
    "answer": "correct option text or letter", // internal; grader will use; can be either option text or letter
    "explanation": "brief explanation of the correct answer"
  }},
  ...
]

Make sure the JSON is valid. If a question is 'open', set "options": [].

Here is the document:
<<DOC>>
{pdf_text[:30000]}
<<ENDDOC>>

Only produce the JSON array.
"""
    raw = gemini_generate_text(prompt, max_output_tokens=1500)
    # parse JSON from response (model may put stray whitespace)
    try:
        parsed = json.loads(raw)
    except Exception as e:
        # Attempt to recover if the model included ```json or other wrappers
        cleaned = raw.strip()
        # Try to find first '[' and last ']' to extract array
        start = cleaned.find('[')
        end = cleaned.rfind(']')
        if start != -1 and end != -1:
            maybe = cleaned[start:end+1]
            parsed = json.loads(maybe)
        else:
            raise RuntimeError(f"Failed to parse JSON from Gemini response. Raw response:\n{raw}") from e

    questions = []
    for obj in parsed:
        q = Question(
            id=int(obj.get("id")),
            qtype=obj.get("type"),
            question=obj.get("question"),
            options=obj.get("options", []) or [],
            correct_answer=str(obj.get("answer")) if obj.get("answer") is not None else None,
            explanation=obj.get("explanation")
        )
        questions.append(q)
    return questions

def ask_gemini_to_grade(pdf_text: str, questions: List[Question], user_answers: Dict[int, str]) -> Dict[str, Any]:
    """
    Ask Gemini to grade the answers. Request strict JSON with per-question feedback and a final score.
    Returns a dict: {"total_score": X, "max_score": Y, "per_question":[{...}, ...]}
    """
    # Build a compact representation
    qlist = []
    for q in questions:
        q_obj = {
            "id": q.id,
            "type": q.qtype,
            "question": q.question,
            "options": q.options,
            "correct_answer": q.correct_answer
        }
        qlist.append(q_obj)

    prompt = f"""
You are an assistant grader. You will be given:
- the document text (delimited by <<DOC>> <<ENDDOC>>),
- the list of questions with their correct answers (delimited by <<QUESTIONS>> <<ENDQUESTIONS>>),
- and the student's answers (delimited by <<ANSWERS>> <<ENDANSWERS>>).

For each question, produce:
- a score (0 or 1 for MCQ and for open questions use partial credit between 0 and 1, to two decimals),
- an explanation of why the answer is correct/incorrect,
- and, for open questions, a short ideal answer.

Then also produce a "total_score" (sum of the per-question scores) and "max_score" (equal to the number of questions, e.g. 6).

Output MUST be valid JSON with this structure:
{{
  "total_score": float,
  "max_score": float,
  "per_question": [
    {{
      "id": 1,
      "score": float,
      "feedback": "text",
      "ideal_answer": "..."  // include for open questions; for MCQ can echo the correct option
    }},
    ...
  ]
}}

Document:
<<DOC>>
{pdf_text[:30000]}
<<ENDDOC>>

Questions+answers:
<<QUESTIONS>>
{json.dumps(qlist)}
<<ENDQUESTIONS>>

Student answers:
<<ANSWERS>>
{json.dumps(user_answers)}
<<ENDANSWERS>>

Produce only the JSON.
"""
    raw = gemini_generate_text(prompt, max_output_tokens=1500)
    try:
        parsed = json.loads(raw)
    except Exception as e:
        cleaned = raw.strip()
        start = cleaned.find('{')
        end = cleaned.rfind('}')
        if start != -1 and end != -1:
            maybe = cleaned[start:end+1]
            parsed = json.loads(maybe)
        else:
            raise RuntimeError(f"Failed to parse grading JSON. Raw response:\n{raw}") from e
    return parsed

# ---------- Streamlit UI ----------
st.set_page_config(page_title="PDF Quiz Generator & Grader", layout="wide")

st.title("PDF Quiz Generator & Grader")

if "questions" not in st.session_state:
    st.session_state["questions"] = None
if "pdf_text" not in st.session_state:
    st.session_state["pdf_text"] = None
if "generated_from_filename" not in st.session_state:
    st.session_state["generated_from_filename"] = None

with st.sidebar:
    st.header("Settings")
    num_questions = st.number_input("Number of questions to generate", min_value=1, max_value=20, value=6, step=1)
    model_choice = st.selectbox("Gemini model (if available)", options=["text-bison-001"], index=0)
    regenerate = st.button("Force regenerate questions (clears previous)")

uploaded_file = st.file_uploader("Upload PDF", type=["pdf"])

if uploaded_file is not None:
    try:
        text = extract_text_from_pdf(uploaded_file)
    except Exception as e:
        st.error(f"Failed to extract text from PDF: {e}")
        st.stop()

    st.session_state["pdf_text"] = text
    # If different file than before, clear generated questions
    if st.session_state.get("generated_from_filename") != uploaded_file.name:
        st.session_state["questions"] = None
        st.session_state["generated_from_filename"] = uploaded_file.name

    st.markdown("### Document preview (first 2000 chars)")
    st.text(text[:2000] + ("..." if len(text) > 2000 else ""))

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate Questions"):
            with st.spinner("Generating questions from the document via Gemini..."):
                try:
                    questions = ask_gemini_for_questions(st.session_state["pdf_text"], num_questions=num_questions)
                    st.session_state["questions"] = questions
                    st.success(f"Generated {len(questions)} questions.")
                except Exception as e:
                    st.error(f"Failed to generate questions: {e}")

    with col2:
        if st.session_state["questions"]:
            if st.button("Clear generated questions"):
                st.session_state["questions"] = None
                st.success("Cleared saved questions.")

# Show questions if present
if st.session_state.get("questions"):
    st.header("Answer the questions")
    answers = {}
    for q in st.session_state["questions"]:
        st.markdown(f"**Q{q.id}.** {q.question}")
        if q.qtype.lower().startswith("mc"):
            # display options
            # label options as A,B,C,D for display
            labels = ["A", "B", "C", "D", "E", "F"]
            opt_map = {labels[i]: opt for i, opt in enumerate(q.options)}
            choices_display = [f"{labels[i]}. {opt}" for i, opt in enumerate(q.options)]
            ans = st.radio(f"Select answer for Q{q.id}", choices_display, key=f"q{q.id}")
            # store just the chosen option text
            selected_label = ans.split(".")[0]
            answers[q.id] = opt_map[selected_label]
        else:
            ans = st.text_input(f"Your answer for Q{q.id}", key=f"q{q.id}")
            answers[q.id] = ans

    if st.button("Submit answers for evaluation"):
        with st.spinner("Evaluating answers with Gemini..."):
            try:
                grading = ask_gemini_to_grade(st.session_state["pdf_text"], st.session_state["questions"], answers)
                st.session_state["grading"] = grading
                st.success("Grading complete.")
            except Exception as e:
                st.error(f"Grading failed: {e}")

# Show grading results if present
if st.session_state.get("grading"):
    grading = st.session_state["grading"]
    st.header("Results")
    st.metric("Score", f"{grading.get('total_score')} / {grading.get('max_score')}")
    for item in grading.get("per_question", []):
        qid = item.get("id")
        score = item.get("score")
        feedback = item.get("feedback")
        ideal = item.get("ideal_answer", "")
        st.subheader(f"Q{qid} — Score: {score}")
        st.write(feedback)
        if ideal:
            st.markdown(f"**Ideal / model answer:** {ideal}")

    # Option: save report to docx
    if st.button("Save results to DOCX"):
        doc = Document()
        doc.add_heading("PDF Quiz Results", level=1)
        doc.add_paragraph(f"Source file: {st.session_state.get('generated_from_filename')}")
        doc.add_paragraph(f"Score: {grading.get('total_score')} / {grading.get('max_score')}")
        doc.add_paragraph("")
        for item in grading.get("per_question", []):
            qid = item.get("id")
            qobj = next((q for q in st.session_state["questions"] if q.id == qid), None)
            if qobj:
                doc.add_heading(f"Q{qid}: {qobj.question}", level=2)
            doc.add_paragraph(f"Score: {item.get('score')}")
            doc.add_paragraph(f"Feedback: {item.get('feedback')}")
            ideal = item.get("ideal_answer", "")
            if ideal:
                doc.add_paragraph(f"Ideal answer: {ideal}")
            doc.add_paragraph("")

        filename = f"quiz_results_{st.session_state.get('generated_from_filename','doc')}.docx"
        doc.save(filename)
        with open(filename, "rb") as f:
            st.download_button("Download results DOCX", f, file_name=filename)
