import fitz
import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
import os
from docx import Document
from io import BytesIO

load_dotenv()
gemini_api_key = os.getenv('GEMINI_API_KEY')

genai.configure(api_key=gemini_api_key)

st.title("PDF Summarizer")
uploaded_file = st.file_uploader("Upload a PDF", type=["pdf"])


# Extract text from PDF
def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text


# Split text into chunks
def chunk_text(text, max_chars=3000):
    return [text[i:i + max_chars] for i in range(0, len(text), max_chars)]


# Summarize using Gemini
def summarize_text(text):
    model = genai.GenerativeModel(st.session_state.gemini_model)
    prompt = f"Summarize the following text into clear bullet points:\n\n{text}"
    response = model.generate_content(prompt)
    return response.text


def create_docx(summary_text):
    doc = Document()
    doc.add_heading('PDF summary', level=1)
    for line in summary_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


if "gemini_model" not in st.session_state:
    st.session_state.gemini_model = "gemini-2.5-flash"

if uploaded_file is not None:
    st.info("Extracting text from PDF...")
    text = extract_text_from_pdf(uploaded_file)

    chunks = chunk_text(text)
    summaries = []

    progress = st.progress(0)
    for i, chunk in enumerate(chunks):
        summary = summarize_text(chunk)
        summaries.append(summary)
        progress.progress((i + 1) / len(chunks))

    final_summary = "\n".join(summaries)

    st.subheader("Summary of the PDF")
    st.write(final_summary)

    # Optional: Download button
    filename = st.text_input("Enter name to save file :", "summary")

    if st.button("Generate Download Links"):
        if filename.strip() == "":
            st.warning("Please enter a valid file name!")
        else:
            st.download_button(
                label="Download as TXT",
                data=final_summary,
                file_name=f"{filename}.txt",
                mime="text/plain"
            )

            docx_buffer = create_docx(final_summary)
            st.download_button(
                label="Download as DOCX",
                data=docx_buffer,
                file_name=f"{filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
